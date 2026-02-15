"""
Générateur de documents Word (.docx) professionnels.
Utilise python-docx pour créer des documents formatés à partir des données entreprise
et du contenu généré par le LLM.
"""

import io
import logging
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Coroutine

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.entreprise import Entreprise
from app.models.fonds_vert import FondsVert
from app.reports.generator import (
    UPLOADS_DIR,
    _load_action_plan,
    _load_carbon,
    _load_credit_score,
    _load_entreprise,
    _load_latest_esg_score,
)

logger = logging.getLogger(__name__)

# Couleurs cohérentes avec l'application (emerald)
EMERALD = RGBColor(0x05, 0x96, 0x69)
EMERALD_LIGHT = RGBColor(0xD1, 0xFA, 0xE5)
DARK_GRAY = RGBColor(0x37, 0x41, 0x51)
MEDIUM_GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)

VALID_TYPES = [
    "lettre_motivation",
    "note_presentation",
    "plan_affaires",
    "engagement_esg",
    "budget_previsionnel",
]

TYPE_LABELS = {
    "lettre_motivation": "Lettre de Motivation",
    "note_presentation": "Note de Présentation",
    "plan_affaires": "Plan d'Affaires Vert",
    "engagement_esg": "Lettre d'Engagement ESG",
    "budget_previsionnel": "Budget Prévisionnel",
}

LlmCallback = Callable[[str], Coroutine[Any, Any, str]]


# ── Helpers de mise en forme ───────────────────────────────────────


def _setup_document(doc: Document, entreprise_nom: str) -> None:
    """Configure les marges, polices par défaut et en-têtes/pieds de page."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # En-tête
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(entreprise_nom)
        run.font.size = Pt(8)
        run.font.color.rgb = MEDIUM_GRAY
        run.font.italic = True

        # Pied de page
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Document confidentiel — Généré par ESG Mefali")
        run.font.size = Pt(7)
        run.font.color.rgb = MEDIUM_GRAY


def _add_title(doc: Document, text: str, level: int = 0) -> None:
    """Ajoute un titre avec la couleur emerald."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = EMERALD if level <= 1 else DARK_GRAY


def _add_section_title(doc: Document, text: str) -> None:
    """Ajoute un titre de section (heading level 2)."""
    _add_title(doc, text, level=2)


def _add_body_text(doc: Document, text: str) -> None:
    """Ajoute du texte corps avec gestion des paragraphes multiples."""
    for paragraph in text.strip().split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(8)


def _add_info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Ajoute un tableau d'informations clé-valeur."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (key, value) in enumerate(rows):
        cell_key = table.rows[i].cells[0]
        cell_val = table.rows[i].cells[1]

        cell_key.text = ""
        run = cell_key.paragraphs[0].add_run(key)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY

        cell_val.text = ""
        run = cell_val.paragraphs[0].add_run(str(value) if value else "—")
        run.font.size = Pt(10)

        # Fond léger sur les lignes paires
        if i % 2 == 0:
            _shade_cell(cell_key, "F3F4F6")
            _shade_cell(cell_val, "F3F4F6")

    # Largeurs des colonnes
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()


def _add_score_table(doc: Document, score: dict) -> None:
    """Ajoute un tableau des scores ESG."""
    headers = ["Pilier", "Score", "Appréciation"]
    piliers = [
        ("Environnement (E)", score.get("score_e", 0)),
        ("Social (S)", score.get("score_s", 0)),
        ("Gouvernance (G)", score.get("score_g", 0)),
        ("Score Global", score.get("score_global", 0)),
    ]

    table = doc.add_table(rows=1 + len(piliers), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # En-tête
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "059669")

    # Données
    for i, (label, val) in enumerate(piliers):
        row = table.rows[i + 1]
        row.cells[0].text = label
        row.cells[1].text = f"{val}/100"
        row.cells[2].text = _appreciation(val)

        if i == len(piliers) - 1:  # Score global en gras
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph()


def _appreciation(score: float) -> str:
    if score >= 80:
        return "Excellent"
    if score >= 60:
        return "Bon"
    if score >= 40:
        return "À améliorer"
    return "Insuffisant"


def _shade_cell(cell, color_hex: str) -> None:
    """Applique une couleur de fond à une cellule."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading.append(shading_elem)


def _add_signature_block(doc: Document, entreprise_nom: str) -> None:
    """Ajoute un bloc de signature."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fait à ________________, le ________________").font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Pour {entreprise_nom}")
    run.font.bold = True
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Nom et qualité du signataire : ________________________________").font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Signature : ").font.size = Pt(10)


async def _load_fonds(db: AsyncSession, fonds_id: str) -> dict | None:
    """Charge les données d'un fonds vert."""
    try:
        result = await db.execute(
            select(FondsVert).where(FondsVert.id == uuid.UUID(fonds_id))
        )
        fonds = result.scalar_one_or_none()
        if not fonds:
            return None
        return {
            "nom": fonds.nom,
            "institution": fonds.institution or "",
            "type": fonds.type or "",
            "montant_min": float(fonds.montant_min) if fonds.montant_min else None,
            "montant_max": float(fonds.montant_max) if fonds.montant_max else None,
            "devise": fonds.devise or "USD",
            "secteurs": fonds.secteurs_json or [],
            "pays_eligibles": fonds.pays_eligibles or [],
            "criteres": fonds.criteres_json or {},
            "date_limite": str(fonds.date_limite) if fonds.date_limite else None,
        }
    except Exception:
        return None


# ── Builders par type de document ──────────────────────────────────


async def _build_lettre_motivation(
    doc: Document, data: dict, llm: LlmCallback
) -> None:
    """Construit une lettre de motivation pour un fonds vert."""
    ent = data["entreprise"]
    fonds = data.get("fonds")
    score = data.get("score")
    instructions = data.get("instructions", "")

    # Date
    now = datetime.now(timezone.utc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(now.strftime("%d/%m/%Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = MEDIUM_GRAY

    # Expéditeur
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(ent["nom"])
    run.font.bold = True
    run.font.size = Pt(12)
    if ent.get("pays"):
        doc.add_paragraph(f"{ent['pays']}")

    doc.add_paragraph()

    # Destinataire
    if fonds:
        p = doc.add_paragraph()
        run = p.add_run(f"À l'attention de : {fonds.get('institution', 'Institution')}")
        run.font.bold = True
        doc.add_paragraph(f"Objet : Candidature au fonds « {fonds['nom']} »")
    else:
        doc.add_paragraph("À l'attention du comité de sélection")
        doc.add_paragraph("Objet : Candidature au financement vert")

    doc.add_paragraph()

    # Corps — LLM
    fonds_desc = ""
    if fonds:
        fonds_desc = (
            f"Le fonds ciblé est « {fonds['nom']} » de {fonds.get('institution', 'N/A')}. "
            f"Montant : {fonds.get('montant_min', '?')} à {fonds.get('montant_max', '?')} {fonds.get('devise', 'USD')}. "
        )

    score_desc = ""
    if score:
        score_desc = (
            f"L'entreprise a obtenu un score ESG de {score.get('score_global', 0)}/100 "
            f"(E={score.get('score_e', 0)}, S={score.get('score_s', 0)}, G={score.get('score_g', 0)}) "
            f"selon le référentiel {score.get('referentiel_nom', 'N/A')}. "
        )

    extra = f"Instructions supplémentaires : {instructions}" if instructions else ""

    prompt = (
        f"Rédige une lettre de motivation professionnelle pour la candidature de "
        f"l'entreprise « {ent['nom']} » (secteur : {ent.get('secteur_activite', 'N/A')}, "
        f"pays : {ent.get('pays', 'N/A')}, effectifs : {ent.get('taille', 'N/A')}). "
        f"{fonds_desc}{score_desc}"
        f"La lettre doit comprendre : 1) Introduction et présentation de l'entreprise, "
        f"2) Motivation et adéquation avec le fonds, 3) Engagements ESG concrets, "
        f"4) Conclusion avec demande d'entretien. "
        f"Ton formel et professionnel, en français. 4-5 paragraphes. "
        f"Ne mets pas d'en-tête ni de date (déjà présents). {extra}"
    )
    body = await llm(prompt)
    _add_body_text(doc, body)

    # Signature
    _add_signature_block(doc, ent["nom"])


async def _build_note_presentation(
    doc: Document, data: dict, llm: LlmCallback
) -> None:
    """Construit une note de présentation de l'entreprise."""
    ent = data["entreprise"]
    score = data.get("score")
    carbon = data.get("carbon")
    instructions = data.get("instructions", "")

    _add_title(doc, "Note de Présentation", level=0)
    p = doc.add_paragraph()
    run = p.add_run(ent["nom"])
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = EMERALD
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Tableau d'informations
    _add_section_title(doc, "1. Informations Générales")
    _add_info_table(doc, [
        ("Raison sociale", ent["nom"]),
        ("Secteur d'activité", ent.get("secteur_activite", "—")),
        ("Pays", ent.get("pays", "—")),
        ("Effectifs", ent.get("taille", "—")),
        ("Chiffre d'affaires", ent.get("chiffre_affaires_formatted", "—")),
        ("Date de création", ent.get("date_creation", "—")),
    ])

    # Activités — LLM
    _add_section_title(doc, "2. Activités et Positionnement")
    extra = f"Instructions : {instructions}" if instructions else ""
    prompt = (
        f"Rédige une présentation des activités et du positionnement de l'entreprise "
        f"« {ent['nom']} » (secteur : {ent.get('secteur_activite', 'N/A')}, "
        f"pays : {ent.get('pays', 'N/A')}, effectifs : {ent.get('taille', 'N/A')}). "
        f"2-3 paragraphes professionnels décrivant les activités principales, "
        f"le positionnement sur le marché et les forces de l'entreprise. {extra}"
    )
    text = await llm(prompt)
    _add_body_text(doc, text)

    # Scores ESG
    if score:
        _add_section_title(doc, "3. Performance ESG")
        p = doc.add_paragraph(f"Référentiel : {score.get('referentiel_nom', 'N/A')}")
        p.runs[0].font.italic = True
        _add_score_table(doc, score)

    # Empreinte carbone
    if carbon:
        _add_section_title(doc, "4. Empreinte Carbone")
        total = carbon.get("total_kg", 0)
        _add_info_table(doc, [
            ("Émissions totales", f"{total:,.0f} kgCO₂e/an"),
            ("Par employé", f"{carbon.get('par_employe', '—')} kgCO₂e"),
        ])

    # Perspectives — LLM
    section_num = 5 if carbon else (4 if score else 3)
    _add_section_title(doc, f"{section_num}. Perspectives et Engagements")
    prompt = (
        f"Rédige les perspectives et engagements futurs de « {ent['nom']} » en matière "
        f"de développement durable et responsabilité ESG. "
        f"Score ESG actuel : {score.get('score_global', 'N/A')}/100. "
        f"2 paragraphes, ton professionnel et engagé. {extra}"
    )
    text = await llm(prompt)
    _add_body_text(doc, text)


async def _build_plan_affaires(
    doc: Document, data: dict, llm: LlmCallback
) -> None:
    """Construit un plan d'affaires vert."""
    ent = data["entreprise"]
    score = data.get("score")
    carbon = data.get("carbon")
    plan = data.get("action_plan")
    credit = data.get("credit_score")
    instructions = data.get("instructions", "")
    extra = f"Instructions : {instructions}" if instructions else ""

    _add_title(doc, "Plan d'Affaires Vert", level=0)
    p = doc.add_paragraph()
    run = p.add_run(ent["nom"])
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = EMERALD
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    now = datetime.now(timezone.utc)
    p = doc.add_paragraph(now.strftime("%B %Y"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = MEDIUM_GRAY

    doc.add_paragraph()

    # 1. Résumé exécutif
    _add_section_title(doc, "1. Résumé Exécutif")
    score_info = f"Score ESG : {score.get('score_global', 'N/A')}/100. " if score else ""
    carbon_info = f"Empreinte carbone : {carbon.get('total_kg', 0):,.0f} kgCO₂e/an. " if carbon else ""
    prompt = (
        f"Rédige le résumé exécutif d'un plan d'affaires vert pour « {ent['nom']} » "
        f"(secteur : {ent.get('secteur_activite', 'N/A')}, pays : {ent.get('pays', 'N/A')}). "
        f"{score_info}{carbon_info}"
        f"2-3 paragraphes présentant la vision verte, les objectifs et les moyens. {extra}"
    )
    _add_body_text(doc, await llm(prompt))

    # 2. Présentation de l'entreprise
    _add_section_title(doc, "2. Présentation de l'Entreprise")
    _add_info_table(doc, [
        ("Raison sociale", ent["nom"]),
        ("Secteur", ent.get("secteur_activite", "—")),
        ("Pays", ent.get("pays", "—")),
        ("Effectifs", ent.get("taille", "—")),
        ("Chiffre d'affaires", ent.get("chiffre_affaires_formatted", "—")),
    ])

    # 3. Analyse du marché vert
    _add_section_title(doc, "3. Analyse du Marché et Opportunités Vertes")
    prompt = (
        f"Analyse le marché vert et les opportunités de développement durable pour "
        f"une entreprise du secteur « {ent.get('secteur_activite', 'N/A')} » "
        f"en {ent.get('pays', 'Afrique')}. "
        f"Tendances du marché, cadre réglementaire, opportunités. 2-3 paragraphes. {extra}"
    )
    _add_body_text(doc, await llm(prompt))

    # 4. Stratégie ESG
    _add_section_title(doc, "4. Stratégie ESG")
    if score:
        _add_score_table(doc, score)

    if plan and plan.get("action_items"):
        doc.add_paragraph("Actions prioritaires :")
        items = plan["action_items"][:8]
        table = doc.add_table(rows=1 + len(items), cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(["Action", "Pilier", "Priorité", "Échéance"]):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "059669")

        for i, item in enumerate(items):
            row = table.rows[i + 1]
            row.cells[0].text = item.get("titre", "")
            row.cells[1].text = item.get("pilier", "").upper() if item.get("pilier") else ""
            row.cells[2].text = item.get("priorite", "")
            row.cells[3].text = item.get("echeance", "—") or "—"
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

        doc.add_paragraph()

    # 5. Plan financier
    _add_section_title(doc, "5. Plan Financier et Budget")
    ca = ent.get("chiffre_affaires_formatted", "N/A")
    prompt = (
        f"Rédige un plan financier et budget prévisionnel pour les investissements verts "
        f"de « {ent['nom']} » (CA : {ca}). "
        f"Inclus : investissements nécessaires, sources de financement visées, "
        f"retour sur investissement attendu, calendrier. 2-3 paragraphes. {extra}"
    )
    _add_body_text(doc, await llm(prompt))

    # 6. Impact
    _add_section_title(doc, "6. Impact Environnemental et Social Attendu")
    prompt = (
        f"Décris l'impact environnemental et social attendu des initiatives vertes de "
        f"« {ent['nom']} » (secteur : {ent.get('secteur_activite', 'N/A')}). "
        f"{carbon_info}"
        f"Inclus : réduction des émissions, emplois verts, impact communautaire. "
        f"2 paragraphes. {extra}"
    )
    _add_body_text(doc, await llm(prompt))


async def _build_engagement_esg(
    doc: Document, data: dict, llm: LlmCallback
) -> None:
    """Construit une lettre d'engagement ESG."""
    ent = data["entreprise"]
    score = data.get("score")
    instructions = data.get("instructions", "")
    extra = f"Instructions : {instructions}" if instructions else ""

    _add_title(doc, "Lettre d'Engagement ESG", level=0)

    doc.add_paragraph()

    # Préambule
    _add_section_title(doc, "Préambule")
    p = doc.add_paragraph(
        f"La société {ent['nom']}, opérant dans le secteur "
        f"{ent.get('secteur_activite', 'N/A')} en {ent.get('pays', 'N/A')}, "
        f"s'engage par la présente à adopter et maintenir des pratiques conformes "
        f"aux standards ESG (Environnement, Social, Gouvernance)."
    )

    doc.add_paragraph()

    # Articles — LLM
    _add_section_title(doc, "Engagements")
    score_info = ""
    if score:
        score_info = (
            f"Score ESG actuel : {score.get('score_global', 0)}/100 "
            f"(E={score.get('score_e', 0)}, S={score.get('score_s', 0)}, G={score.get('score_g', 0)}). "
        )
    prompt = (
        f"Rédige 6 à 8 engagements ESG formels pour « {ent['nom']} » "
        f"(secteur : {ent.get('secteur_activite', 'N/A')}). {score_info}"
        f"Format : chaque engagement commence par 'Article X :' suivi du titre puis "
        f"d'une description de 2-3 lignes. Couvre les 3 piliers E, S et G. "
        f"Ton formel et engageant. {extra}"
    )
    text = await llm(prompt)
    _add_body_text(doc, text)

    # Tableau indicateurs
    _add_section_title(doc, "Indicateurs de Suivi")
    indicators = [
        ("Environnement", "Score E", str(score.get("score_e", "—")) if score else "—", "Annuel"),
        ("Social", "Score S", str(score.get("score_s", "—")) if score else "—", "Annuel"),
        ("Gouvernance", "Score G", str(score.get("score_g", "—")) if score else "—", "Annuel"),
        ("Global", "Score ESG", str(score.get("score_global", "—")) if score else "—", "Annuel"),
    ]
    table = doc.add_table(rows=1 + len(indicators), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(["Pilier", "Indicateur", "Valeur actuelle", "Fréquence"]):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "059669")

    for i, (pilier, ind, val, freq) in enumerate(indicators):
        row = table.rows[i + 1]
        row.cells[0].text = pilier
        row.cells[1].text = ind
        row.cells[2].text = val
        row.cells[3].text = freq

    doc.add_paragraph()

    # Signature
    _add_signature_block(doc, ent["nom"])


async def _build_budget_previsionnel(
    doc: Document, data: dict, llm: LlmCallback
) -> None:
    """Construit un budget prévisionnel pour un projet vert."""
    ent = data["entreprise"]
    fonds = data.get("fonds")
    score = data.get("score")
    instructions = data.get("instructions", "")
    extra = f"Instructions : {instructions}" if instructions else ""

    _add_title(doc, "Budget Prévisionnel", level=0)
    p = doc.add_paragraph()
    run = p.add_run(f"Projet d'investissement vert — {ent['nom']}")
    run.font.size = Pt(12)
    run.font.color.rgb = EMERALD
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Contexte
    _add_section_title(doc, "1. Contexte du Projet")
    fonds_info = ""
    if fonds:
        montant = ""
        if fonds.get("montant_min") and fonds.get("montant_max"):
            montant = f" (montant : {fonds['montant_min']:,.0f} — {fonds['montant_max']:,.0f} {fonds.get('devise', 'USD')})"
        fonds_info = f"Fonds ciblé : {fonds['nom']} de {fonds.get('institution', 'N/A')}{montant}. "

    prompt = (
        f"Rédige le contexte d'un projet d'investissement vert pour « {ent['nom']} » "
        f"(secteur : {ent.get('secteur_activite', 'N/A')}, pays : {ent.get('pays', 'N/A')}, "
        f"CA : {ent.get('chiffre_affaires_formatted', 'N/A')}). {fonds_info}"
        f"2 paragraphes décrivant le projet, ses objectifs et sa justification. {extra}"
    )
    _add_body_text(doc, await llm(prompt))

    # Budget — LLM génère un tableau structuré
    _add_section_title(doc, "2. Budget Détaillé")
    prompt = (
        f"Génère un budget prévisionnel détaillé pour un projet vert de « {ent['nom']} » "
        f"(secteur : {ent.get('secteur_activite', 'N/A')}). {fonds_info}"
        f"Liste 6 à 10 postes budgétaires, chacun sur une ligne avec ce format exact :\n"
        f"POSTE | MONTANT | JUSTIFICATION\n"
        f"Exemple : Panneaux solaires | 15 000 000 XOF | Installation de 20 panneaux\n"
        f"À la fin, ajoute une ligne TOTAL avec le montant total.\n"
        f"Utilise des montants réalistes en XOF. {extra}"
    )
    budget_text = await llm(prompt)

    # Parser le texte LLM en tableau
    lines = [l.strip() for l in budget_text.strip().split("\n") if "|" in l]
    if lines:
        # Vérifier si la première ligne est un en-tête
        first = lines[0].lower()
        if "poste" in first or "montant" in first or "désignation" in first:
            lines = lines[1:]
        # Ignorer les lignes de séparation (---|--)
        lines = [l for l in lines if not all(c in "-| " for c in l)]

        if lines:
            table = doc.add_table(rows=1 + len(lines), cols=3)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(["Poste", "Montant", "Justification"]):
                cell = table.rows[0].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade_cell(cell, "059669")

            for i, line in enumerate(lines):
                parts = [p.strip() for p in line.split("|")]
                row = table.rows[i + 1]
                for j in range(min(3, len(parts))):
                    row.cells[j].text = parts[j]
                    for p in row.cells[j].paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

                # Total en gras
                is_total = parts[0].lower().startswith("total")
                if is_total:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                    _shade_cell(row.cells[0], "D1FAE5")
                    _shade_cell(row.cells[1], "D1FAE5")
                    _shade_cell(row.cells[2], "D1FAE5")

            doc.add_paragraph()

    # Sources de financement
    _add_section_title(doc, "3. Sources de Financement")
    prompt = (
        f"Décris les sources de financement envisagées pour ce projet vert : "
        f"fonds propres, fonds verts, subventions, prêts bancaires verts. {fonds_info}"
        f"Indique la répartition envisagée. 2 paragraphes. {extra}"
    )
    _add_body_text(doc, await llm(prompt))

    # Calendrier
    _add_section_title(doc, "4. Calendrier de Mise en Œuvre")
    prompt = (
        f"Propose un calendrier de mise en œuvre sur 12 à 24 mois pour le projet vert "
        f"de « {ent['nom']} ». Liste 4 à 6 phases avec leurs échéances. "
        f"Format : une ligne par phase, concis. {extra}"
    )
    text = await llm(prompt)
    _add_body_text(doc, text)


# ── Fonction principale ───────────────────────────────────────────


BUILDERS = {
    "lettre_motivation": _build_lettre_motivation,
    "note_presentation": _build_note_presentation,
    "plan_affaires": _build_plan_affaires,
    "engagement_esg": _build_engagement_esg,
    "budget_previsionnel": _build_budget_previsionnel,
}


async def generate_word_document(
    entreprise_id: str,
    document_type: str,
    db: AsyncSession,
    llm_callback: LlmCallback,
    fonds_id: str | None = None,
    instructions: str | None = None,
) -> tuple[bytes, str]:
    """
    Génère un document Word (.docx) professionnel.

    Returns:
        tuple[bytes, str]: (contenu du fichier, nom du fichier)
    """
    if document_type not in VALID_TYPES:
        raise ValueError(
            f"Type de document inconnu : '{document_type}'. "
            f"Types valides : {', '.join(VALID_TYPES)}"
        )

    # Charger les données
    entreprise = await _load_entreprise(db, entreprise_id)
    if not entreprise:
        raise ValueError("Entreprise introuvable")

    score = await _load_latest_esg_score(db, entreprise_id)
    carbon = await _load_carbon(db, entreprise_id)
    credit = await _load_credit_score(db, entreprise_id)
    action_plan = await _load_action_plan(db, entreprise_id)

    fonds = None
    if fonds_id:
        fonds = await _load_fonds(db, fonds_id)

    data = {
        "entreprise": entreprise,
        "score": score,
        "carbon": carbon,
        "credit_score": credit,
        "action_plan": action_plan,
        "fonds": fonds,
        "instructions": instructions or "",
    }

    # Créer le document
    doc = Document()
    _setup_document(doc, entreprise["nom"])

    # Construire le contenu
    builder = BUILDERS[document_type]
    await builder(doc, data, llm_callback)

    # Sauvegarder
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    now = datetime.now(timezone.utc)
    nom_clean = entreprise["nom"].replace(" ", "_").replace("/", "_")
    filename = f"{document_type}_{nom_clean}_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = UPLOADS_DIR / filename

    # Écrire en mémoire puis sur disque
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    filepath.write_bytes(docx_bytes)

    logger.info("Document Word généré : %s (%d Ko)", filename, len(docx_bytes) // 1024)
    return docx_bytes, filename
