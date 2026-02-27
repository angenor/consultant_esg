"""
Template : Note d'Impact ESG dédiée à la candidature.
Document de 3-4 pages analysant la performance ESG de l'entreprise
et son alignement avec le fonds ciblé.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.documents.word_generator import (
    EMERALD,
    DARK_GRAY,
    MEDIUM_GRAY,
    LlmCallback,
    _add_title,
    _add_section_title,
    _add_body_text,
    _add_info_table,
    _add_score_table,
    _add_placeholder,
    _appreciation,
    _shade_cell,
)


async def build_note_impact_esg(
    doc: Document, data: dict, llm: LlmCallback, mode: str = "complet"
) -> None:
    """
    Construit une note d'impact ESG dédiée à la candidature.

    data attendu :
      - entreprise: dict (nom, secteur_activite, pays)
      - fonds: dict (nom, institution) | None
      - score: dict (score_global, score_e, score_s, score_g, referentiel_nom) | None
      - carbon: dict (total_kg, par_employe) | None
      - action_plan: dict (action_items: list) | None
      - instructions: str
    """
    ent = data["entreprise"]
    fonds = data.get("fonds")
    score = data.get("score")
    carbon = data.get("carbon")
    action_plan = data.get("action_plan")
    instructions = data.get("instructions", "")
    extra = f"Instructions : {instructions}" if instructions else ""

    is_template = mode == "template_vierge"

    # Titre
    _add_title(doc, "Note d'Impact ESG", level=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ent["nom"])
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = EMERALD

    if fonds:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Candidature au fonds : {fonds['nom']}")
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = MEDIUM_GRAY

    doc.add_paragraph()

    # 1. Score ESG actuel
    _add_section_title(doc, "1. Performance ESG Actuelle")

    if is_template:
        _add_placeholder(doc, "Indiquez votre score ESG global et par pilier (E, S, G). Si vous n'avez pas de score, décrivez vos pratiques ESG actuelles.")
        # Tableau vierge
        headers = ["Pilier", "Score (/100)", "Appréciation"]
        table = doc.add_table(rows=5, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "059669")

        piliers = ["Environnement (E)", "Social (S)", "Gouvernance (G)", "Score Global"]
        for i, pilier in enumerate(piliers):
            row = table.rows[i + 1]
            row.cells[0].text = pilier
            row.cells[1].text = "[__/100]"
            row.cells[2].text = "[À compléter]"
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

        doc.add_paragraph()
    elif score:
        p = doc.add_paragraph(f"Référentiel : {score.get('referentiel_nom', 'N/A')}")
        p.runs[0].font.italic = True
        _add_score_table(doc, score)
    else:
        doc.add_paragraph("Aucun score ESG disponible. L'évaluation est recommandée avant la candidature.")

    # 2. Analyse par pilier
    _add_section_title(doc, "2. Analyse par Pilier ESG")

    if is_template:
        for pilier in ["Environnement", "Social", "Gouvernance"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{pilier}")
            run.font.bold = True
            run.font.size = Pt(11)
            _add_placeholder(doc, f"Décrivez vos points forts et axes d'amélioration pour le pilier {pilier}.")
    else:
        score_detail = ""
        if score:
            score_detail = (
                f"Scores actuels : E={score.get('score_e', 0)}/100, "
                f"S={score.get('score_s', 0)}/100, G={score.get('score_g', 0)}/100. "
            )
        prompt = (
            f"Analyse la performance ESG de « {ent['nom']} » par pilier "
            f"(Environnement, Social, Gouvernance). {score_detail}"
            f"Pour chaque pilier, identifie les points forts et les axes d'amélioration. "
            f"Secteur : {ent.get('secteur_activite', 'N/A')}, pays : {ent.get('pays', 'N/A')}. "
            f"3 sous-sections (E, S, G) avec 2-3 paragraphes chacune. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # 3. Alignement avec le fonds cible
    _add_section_title(doc, "3. Alignement avec le Référentiel du Fonds")

    if is_template:
        _add_placeholder(doc, "Expliquez comment votre entreprise et votre projet s'alignent avec les critères et priorités du fonds ciblé.")
    else:
        fonds_info = ""
        if fonds:
            fonds_info = (
                f"Le fonds ciblé est « {fonds['nom']} » de {fonds.get('institution', 'N/A')}. "
            )
        prompt = (
            f"Analyse l'alignement de « {ent['nom']} » avec le référentiel du fonds cible. "
            f"{fonds_info}"
            f"Secteur : {ent.get('secteur_activite', 'N/A')}. "
            f"Comment l'entreprise répond-elle aux critères du fonds ? "
            f"2-3 paragraphes. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # 4. Plan d'amélioration ESG
    _add_section_title(doc, "4. Plan d'Amélioration ESG")

    if is_template:
        _add_placeholder(doc, "Décrivez votre plan d'amélioration ESG : actions prioritaires, échéances, responsables, budget dédié.")

        headers = ["Action", "Pilier", "Échéance", "Budget"]
        table = doc.add_table(rows=5, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "059669")

        for i in range(1, 5):
            for j in range(4):
                table.rows[i].cells[j].text = ""
                run = table.rows[i].cells[j].paragraphs[0].add_run("[À compléter]")
                run.font.size = Pt(9)
                run.font.color.rgb = MEDIUM_GRAY

        doc.add_paragraph()
    else:
        # Utiliser le plan d'action existant s'il y en a un
        if action_plan and action_plan.get("action_items"):
            items = action_plan["action_items"][:6]
            table = doc.add_table(rows=1 + len(items), cols=4)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(["Action", "Pilier", "Priorité", "Échéance"]):
                cell = table.rows[0].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade_cell(cell, "059669")

            for i, item in enumerate(items):
                row = table.rows[i + 1]
                row.cells[0].text = item.get("titre", "")
                row.cells[1].text = (item.get("pilier") or "").upper()
                row.cells[2].text = item.get("priorite", "")
                row.cells[3].text = item.get("echeance", "—") or "—"
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

            doc.add_paragraph()

        score_info = f"Score ESG actuel : {score.get('score_global', 'N/A')}/100. " if score else ""
        prompt = (
            f"Propose un plan d'amélioration ESG pour « {ent['nom']} » lié au projet financé. "
            f"{score_info}"
            f"Actions concrètes, échéances, indicateurs de suivi. 2-3 paragraphes. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # 5. Benchmarks sectoriels
    _add_section_title(doc, "5. Positionnement Sectoriel")

    if is_template:
        _add_placeholder(doc, "Comparez votre performance ESG avec la moyenne de votre secteur et les meilleures pratiques.")
    else:
        score_bench = f"Score global : {score.get('score_global', 'N/A')}/100. " if score else ""
        prompt = (
            f"Compare la performance ESG de « {ent['nom']} » avec les standards "
            f"du secteur « {ent.get('secteur_activite', 'N/A')} » en {ent.get('pays', 'Afrique')}. "
            f"{score_bench}"
            f"Comment se positionne l'entreprise ? Points de différenciation. "
            f"2 paragraphes. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # 6. Indicateurs de suivi
    _add_section_title(doc, "6. Indicateurs de Suivi Proposés")

    if is_template:
        _add_placeholder(doc, "Proposez 5-8 indicateurs de suivi ESG avec fréquence de mesure et cibles.")
    else:
        # Tableau d'indicateurs
        indicators = [
            ("Score ESG global", "Annuel", score.get("score_global", "—") if score else "—", "+10% / an"),
            ("Score Environnement", "Annuel", score.get("score_e", "—") if score else "—", "+15% / an"),
            ("Score Social", "Annuel", score.get("score_s", "—") if score else "—", "+10% / an"),
            ("Score Gouvernance", "Annuel", score.get("score_g", "—") if score else "—", "+10% / an"),
            ("Émissions CO₂", "Annuel", f"{carbon.get('total_kg', 0):,.0f} kgCO₂e" if carbon else "—", "-20% / 3 ans"),
            ("Emplois verts créés", "Semestriel", "—", "À définir"),
        ]

        table = doc.add_table(rows=1 + len(indicators), cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(["Indicateur", "Fréquence", "Valeur actuelle", "Cible"]):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "059669")

        for i, (ind, freq, val, cible) in enumerate(indicators):
            row = table.rows[i + 1]
            row.cells[0].text = ind
            row.cells[1].text = freq
            row.cells[2].text = str(val)
            row.cells[3].text = cible
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

        doc.add_paragraph()
