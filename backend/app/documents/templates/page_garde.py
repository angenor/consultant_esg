"""
Template : Page de garde du dossier de candidature.
Génère une page de couverture professionnelle (pas d'appel LLM).
"""

from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.documents.word_generator import (
    EMERALD,
    DARK_GRAY,
    MEDIUM_GRAY,
    _setup_document,
    _shade_cell,
)


def build_page_garde(doc: Document, data: dict) -> None:
    """
    Construit la page de garde du dossier de candidature.

    data attendu :
      - entreprise: dict (nom, secteur_activite, pays)
      - fonds: dict (nom, institution) | None
      - intermediaire_nom: str | None
      - documents_inclus: list[str]  (noms des documents dans le dossier)
      - reference: str  (référence auto-générée)
    """
    ent = data["entreprise"]
    fonds = data.get("fonds")
    intermediaire_nom = data.get("intermediaire_nom")
    documents_inclus = data.get("documents_inclus", [])
    reference = data.get("reference", "")
    now = datetime.now(timezone.utc)

    # Espace haut
    for _ in range(3):
        doc.add_paragraph()

    # Titre principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOSSIER DE CANDIDATURE")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = EMERALD

    # Sous-titre fonds
    if fonds:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(fonds["nom"])
        run.font.size = Pt(18)
        run.font.color.rgb = DARK_GRAY

        if fonds.get("institution"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(fonds["institution"])
            run.font.size = Pt(14)
            run.font.color.rgb = MEDIUM_GRAY

    doc.add_paragraph()

    # Ligne séparatrice (simulée par un paragraphe avec bordure)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    run.font.color.rgb = EMERALD
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Entreprise
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Présenté par")
    run.font.size = Pt(11)
    run.font.color.rgb = MEDIUM_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ent["nom"])
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = DARK_GRAY

    if ent.get("secteur_activite"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ent["secteur_activite"])
        run.font.size = Pt(12)
        run.font.color.rgb = MEDIUM_GRAY

    if ent.get("pays"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ent["pays"])
        run.font.size = Pt(12)
        run.font.color.rgb = MEDIUM_GRAY

    # Intermédiaire
    if intermediaire_nom:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Via : {intermediaire_nom}")
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    # Infos clés
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = 1  # CENTER

    rows_data = [
        ("Date de soumission", now.strftime("%d/%m/%Y")),
        ("Référence", reference),
        ("Confidentialité", "Document confidentiel"),
    ]
    for i, (label, value) in enumerate(rows_data):
        cell_l = info_table.rows[i].cells[0]
        cell_r = info_table.rows[i].cells[1]
        cell_l.text = ""
        cell_r.text = ""
        run = cell_l.paragraphs[0].add_run(label)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        run = cell_r.paragraphs[0].add_run(value)
        run.font.size = Pt(10)
        cell_l.width = Cm(5)
        cell_r.width = Cm(8)

    doc.add_paragraph()

    # Table des matières
    if documents_inclus:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Table des matières")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = EMERALD

        doc.add_paragraph()

        toc_table = doc.add_table(rows=len(documents_inclus), cols=2)
        toc_table.alignment = 1
        for i, doc_name in enumerate(documents_inclus):
            cell_num = toc_table.rows[i].cells[0]
            cell_name = toc_table.rows[i].cells[1]
            cell_num.text = ""
            cell_name.text = ""
            run = cell_num.paragraphs[0].add_run(f"{i + 1:02d}")
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = EMERALD
            run = cell_name.paragraphs[0].add_run(doc_name)
            run.font.size = Pt(10)
            cell_num.width = Cm(1.5)
            cell_name.width = Cm(12)

    # Pied de page
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document confidentiel — Généré par ESG Mefali")
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.italic = True
