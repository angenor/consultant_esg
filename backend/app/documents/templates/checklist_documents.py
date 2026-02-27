"""
Template : Checklist des documents du dossier de candidature.
Affiche un tableau récapitulatif avec le statut de chaque document (pas d'appel LLM).
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.documents.word_generator import (
    EMERALD,
    DARK_GRAY,
    MEDIUM_GRAY,
    _add_title,
    _shade_cell,
)

# Couleurs de statut
_STATUS_COLORS = {
    "inclus": "D1FAE5",       # vert clair
    "a_fournir": "FEF3C7",   # jaune clair
    "non_applicable": "F3F4F6",  # gris clair
}

_STATUS_LABELS = {
    "inclus": "Inclus",
    "a_fournir": "À fournir",
    "non_applicable": "Non applicable",
}


def build_checklist(doc: Document, data: dict) -> None:
    """
    Construit la checklist des documents du dossier.

    data attendu :
      - fonds: dict (nom) | None
      - intermediaire_nom: str | None
      - documents_status: list[dict]
            Chaque dict : {nom: str, statut: str, notes: str}
            statut parmi : "inclus", "a_fournir", "non_applicable"
    """
    fonds = data.get("fonds")
    intermediaire_nom = data.get("intermediaire_nom")
    documents_status = data.get("documents_status", [])

    _add_title(doc, "Checklist des Documents", level=0)

    # Sous-titre
    subtitle_parts = []
    if fonds:
        subtitle_parts.append(f"Fonds : {fonds['nom']}")
    if intermediaire_nom:
        subtitle_parts.append(f"Intermédiaire : {intermediaire_nom}")
    if subtitle_parts:
        p = doc.add_paragraph(" — ".join(subtitle_parts))
        p.runs[0].font.color.rgb = MEDIUM_GRAY
        p.runs[0].font.italic = True

    doc.add_paragraph()

    if not documents_status:
        doc.add_paragraph("Aucun document dans le dossier.")
        return

    # Tableau
    headers = ["#", "Document", "Statut", "Notes"]
    table = doc.add_table(rows=1 + len(documents_status), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tête
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "059669")

    # Lignes
    for i, doc_info in enumerate(documents_status):
        row = table.rows[i + 1]
        statut = doc_info.get("statut", "a_fournir")
        color = _STATUS_COLORS.get(statut, "F3F4F6")

        row.cells[0].text = ""
        run = row.cells[0].paragraphs[0].add_run(str(i + 1))
        run.font.size = Pt(9)
        run.font.bold = True

        row.cells[1].text = ""
        run = row.cells[1].paragraphs[0].add_run(doc_info.get("nom", ""))
        run.font.size = Pt(9)

        row.cells[2].text = ""
        label = _STATUS_LABELS.get(statut, statut)
        run = row.cells[2].paragraphs[0].add_run(label)
        run.font.size = Pt(9)
        run.font.bold = True
        _shade_cell(row.cells[2], color)

        row.cells[3].text = ""
        run = row.cells[3].paragraphs[0].add_run(doc_info.get("notes", ""))
        run.font.size = Pt(9)
        run.font.color.rgb = MEDIUM_GRAY

    # Largeurs
    for row in table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(6)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(6)

    doc.add_paragraph()

    # Légende
    p = doc.add_paragraph()
    run = p.add_run("Légende : ")
    run.font.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_GRAY

    for statut, label in _STATUS_LABELS.items():
        run = p.add_run(f"  {label}")
        run.font.size = Pt(8)
        run.font.color.rgb = MEDIUM_GRAY
