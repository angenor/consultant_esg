"""
Template : Fiche Projet Vert.
Document de 2-3 pages décrivant le projet d'investissement vert,
son contexte, ses objectifs et son impact attendu.
"""

from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.documents.word_generator import (
    EMERALD,
    DARK_GRAY,
    MEDIUM_GRAY,
    LlmCallback,
    _add_title,
    _add_section_title,
    _add_body_text,
    _add_info_table,
    _add_score_table,
    _add_placeholder,
    _shade_cell,
    get_fonds_prompt,
)


async def build_fiche_projet(
    doc: Document, data: dict, llm: LlmCallback, mode: str = "complet"
) -> None:
    """
    Construit une fiche projet vert.

    data attendu :
      - entreprise: dict (nom, secteur_activite, pays, taille, chiffre_affaires_formatted)
      - fonds: dict (nom, institution, montant_min, montant_max, devise) | None
      - score: dict (score_global, score_e, score_s, score_g) | None
      - carbon: dict (total_kg, par_employe) | None
      - intermediaire_nom: str | None
      - instructions: str
    """
    ent = data["entreprise"]
    fonds = data.get("fonds")
    score = data.get("score")
    carbon = data.get("carbon")
    intermediaire_nom = data.get("intermediaire_nom")
    instructions = data.get("instructions", "")
    extra = f"Instructions : {instructions}" if instructions else ""
    now = datetime.now(timezone.utc)

    is_template = mode == "template_vierge"

    # En-tête
    _add_title(doc, "Fiche Projet Vert", level=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = [ent["nom"]]
    if fonds:
        parts.append(f"— {fonds['nom']}")
    run = p.add_run(" ".join(parts))
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = EMERALD

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(now.strftime("%B %Y"))
    run.font.color.rgb = MEDIUM_GRAY

    if intermediaire_nom:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Via : {intermediaire_nom}")
        run.font.italic = True
        run.font.color.rgb = MEDIUM_GRAY

    doc.add_paragraph()

    # 1. Résumé exécutif
    _add_section_title(doc, "1. Résumé Exécutif")
    if is_template:
        _add_placeholder(doc, "Résumez le projet en 5-10 lignes : objectifs, bénéficiaires, montant demandé, impact attendu.")
    else:
        fonds_desc = ""
        if fonds:
            fonds_desc = (
                f"Le fonds ciblé est « {fonds['nom']} » de {fonds.get('institution', 'N/A')}. "
            )
        prompt = (
            f"Rédige un résumé exécutif pour la fiche projet vert de « {ent['nom']} » "
            f"(secteur : {ent.get('secteur_activite', 'N/A')}, pays : {ent.get('pays', 'N/A')}). "
            f"{fonds_desc}"
            f"3-4 paragraphes présentant le projet, ses objectifs et les résultats attendus. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # 2. Contexte et problématique
    _add_section_title(doc, "2. Contexte et Problématique")
    if is_template:
        _add_placeholder(doc, "Décrivez le contexte sectoriel, les défis environnementaux et sociaux identifiés, et la problématique à laquelle le projet répond.")
    else:
        prompt = (
            f"Décris le contexte et la problématique justifiant un projet vert pour "
            f"« {ent['nom']} » (secteur : {ent.get('secteur_activite', 'N/A')}, "
            f"pays : {ent.get('pays', 'N/A')}). "
            f"Inclus les défis environnementaux et sociaux du secteur. 2-3 paragraphes. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # 3. Description du projet
    _add_section_title(doc, "3. Description du Projet")

    if is_template:
        subsections = [
            ("Objectifs", "Listez les objectifs principaux du projet (3-5 objectifs)."),
            ("Activités prévues", "Décrivez les activités principales qui seront menées."),
            ("Zone géographique", "Précisez la zone d'intervention du projet."),
            ("Bénéficiaires", "Identifiez les bénéficiaires directs et indirects."),
        ]
        for title, instruction in subsections:
            p = doc.add_paragraph()
            run = p.add_run(f"{title}")
            run.font.bold = True
            run.font.size = Pt(11)
            _add_placeholder(doc, instruction)
    else:
        prompt = (
            f"Décris un projet d'investissement vert pour « {ent['nom']} » "
            f"(secteur : {ent.get('secteur_activite', 'N/A')}). "
            f"Structure la réponse en 4 sous-parties : "
            f"Objectifs (liste 3-5), Activités prévues, Zone géographique, Bénéficiaires. "
            f"2-3 paragraphes par sous-partie. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # 4. Impact environnemental et social attendu
    _add_section_title(doc, "4. Impact Environnemental et Social Attendu")

    if is_template:
        impact_items = [
            ("Réduction CO₂ estimée", "Estimez la réduction des émissions de GES en tCO₂e/an."),
            ("Emplois créés/préservés", "Indiquez le nombre d'emplois directs et indirects."),
            ("Contribution aux ODD", "Listez les ODD auxquels le projet contribue."),
        ]
        for title, instruction in impact_items:
            p = doc.add_paragraph()
            run = p.add_run(f"{title}")
            run.font.bold = True
            run.font.size = Pt(11)
            _add_placeholder(doc, instruction)
    else:
        carbon_info = ""
        if carbon:
            carbon_info = f"Empreinte carbone actuelle : {carbon.get('total_kg', 0):,.0f} kgCO₂e/an. "
        score_info = ""
        if score:
            score_info = f"Score ESG : {score.get('score_global', 0)}/100. "

        prompt = (
            f"Décris l'impact environnemental et social attendu du projet vert de "
            f"« {ent['nom']} ». {carbon_info}{score_info}"
            f"Inclus : réduction des émissions (en tCO₂e), emplois créés, "
            f"contribution aux ODD pertinents. 2-3 paragraphes. {extra}"
        )
        _add_body_text(doc, await llm(prompt))

    # Tableau score ESG si disponible
    if score and not is_template:
        _add_score_table(doc, score)

    # 5. Budget résumé
    _add_section_title(doc, "5. Budget Résumé")

    if is_template:
        # Tableau vierge
        headers = ["Poste budgétaire", "Montant", "Source de financement"]
        table = doc.add_table(rows=6, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "059669")

        for i in range(1, 5):
            for j in range(3):
                table.rows[i].cells[j].text = ""
                run = table.rows[i].cells[j].paragraphs[0].add_run("[À compléter]")
                run.font.size = Pt(9)
                run.font.color.rgb = MEDIUM_GRAY

        # Ligne total
        row_total = table.rows[5]
        row_total.cells[0].text = ""
        run = row_total.cells[0].paragraphs[0].add_run("TOTAL")
        run.font.bold = True
        run.font.size = Pt(9)
        for j in range(3):
            _shade_cell(row_total.cells[j], "D1FAE5")

        doc.add_paragraph()
    else:
        fonds_info = ""
        if fonds:
            fonds_info = f"Fonds ciblé : {fonds['nom']}. "
        prompt = (
            f"Propose un budget résumé (4-6 postes) pour le projet vert de « {ent['nom']} ». "
            f"{fonds_info}Utilise le format : POSTE | MONTANT (XOF) | SOURCE. "
            f"Ajoute une ligne TOTAL. Montants réalistes. {extra}"
        )
        budget_text = await llm(prompt)
        lines = [l.strip() for l in budget_text.strip().split("\n") if "|" in l]
        if lines:
            first = lines[0].lower()
            if "poste" in first or "montant" in first:
                lines = lines[1:]
            lines = [l for l in lines if not all(c in "-| " for c in l)]

            if lines:
                table = doc.add_table(rows=1 + len(lines), cols=3)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for j, h in enumerate(["Poste", "Montant", "Source"]):
                    cell = table.rows[0].cells[j]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(h)
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _shade_cell(cell, "059669")

                for i, line in enumerate(lines):
                    parts = [p.strip() for p in line.split("|")]
                    row = table.rows[i + 1]
                    for j in range(min(3, len(parts))):
                        row.cells[j].text = parts[j]
                        for p in row.cells[j].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)
                    if parts[0].lower().startswith("total"):
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.bold = True
                            _shade_cell(cell, "D1FAE5")

                doc.add_paragraph()

    # 6. Calendrier de mise en œuvre
    _add_section_title(doc, "6. Calendrier de Mise en Œuvre")

    if is_template:
        headers = ["Phase", "Activités", "Période", "Responsable"]
        table = doc.add_table(rows=5, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "059669")

        for i in range(1, 5):
            row = table.rows[i]
            row.cells[0].text = ""
            run = row.cells[0].paragraphs[0].add_run(f"Phase {i}")
            run.font.size = Pt(9)
            run.font.bold = True
            for j in range(1, 4):
                row.cells[j].text = ""
                run = row.cells[j].paragraphs[0].add_run("[À compléter]")
                run.font.size = Pt(9)
                run.font.color.rgb = MEDIUM_GRAY

        doc.add_paragraph()
    else:
        prompt = (
            f"Propose un calendrier de mise en œuvre sur 12-24 mois pour le projet vert "
            f"de « {ent['nom']} ». Liste 4-6 phases avec périodes et activités clés. "
            f"Format concis. {extra}"
        )
        _add_body_text(doc, await llm(prompt))
